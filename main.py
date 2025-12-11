#!/usr/bin/env python
"""Command line interface for running Qwen3-VL-8B OCR over scanned books."""

from __future__ import annotations

import argparse
import json
import logging
import os
import re
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import torch
from PIL import Image
from docx import Document
from docx.shared import Inches
from qwen_vl_utils import process_vision_info
from tqdm.auto import tqdm
from transformers import Qwen3VLForConditionalGeneration, AutoProcessor

LOGGER = logging.getLogger("qwen_ocr")

PAGE_PROMPT = textwrap.dedent(
    """
    Bạn là trợ lý OCR tiếng Việt. Hãy đọc toàn bộ hình trang sách và tạo bản phục dựng Markdown bảo tồn tiêu đề, chú thích, bảng và công thức. Để giữ nguyên vị trí biểu đồ/hình ảnh, hãy trả về JSON với cấu trúc:
    {
      "markdown": "...Markdown sử dụng placeholder {{ASSET:<id>}} cho từng hình...",
      "assets": [
        {
          "id": "fig-1",
          "caption": "Chú thích ngắn gọn",
          "type": "figure|chart|formula",
          "bbox": [x1, y1, x2, y2]
        }
      ]
    }
    - Markdown phải ở đúng trình tự đọc từ trái sang phải, trên xuống dưới.
    - Giữ nguyên dấu tiếng Việt và định dạng (## cấp 2 cho tiêu đề cấp trang, ### cho khối con, danh sách với - hoặc 1.).
    - Nếu không tìm thấy tài liệu, trả về markdown rỗng và assets rỗng.
    - Không được thêm lời giải thích bên ngoài JSON.
    """
).strip()

ASSET_PLACEHOLDER_RE = re.compile(r"\{\{ASSET:([A-Za-z0-9_-]+)\}\}")
IMAGE_LINE_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)$")
IMAGE_INLINE_RE = re.compile(r"!\[(.*?)\]\((.*?)\)")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
NUMBERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="OCR Vietnamese medicine books with Qwen3-VL-8B")
    parser.add_argument("--source", default="docs", type=Path, help="Directory containing book folders")
    parser.add_argument("--output", default="outputs", type=Path, help="Directory for OCR results")
    parser.add_argument(
        "--books",
        nargs="*",
        default=None,
        help="Names of book folders to process (default: all folders under --source)",
    )
    parser.add_argument(
        "--model-id",
        default="Qwen/Qwen3-VL-8B-Instruct",
        help="Hugging Face model repository to load",
    )
    parser.add_argument("--max-new-tokens", type=int, default=1024, help="Max tokens to generate per page")
    parser.add_argument(
        "--temperature",
        type=float,
        default=0.0,
        help="Temperature > 0 enables sampling; 0 switches to greedy decoding",
    )
    parser.add_argument(
        "--dtype",
        choices=["auto", "float16", "bfloat16", "float32"],
        default="auto",
        help="Torch dtype override",
    )
    parser.add_argument("--device-map", default="auto", help="transformers device_map setting")
    parser.add_argument("--min-pixels", type=int, default=512 * 512, help="Min pixels per side for processor")
    parser.add_argument("--max-pixels", type=int, default=2048 * 2048, help="Max pixels per side for processor")
    parser.add_argument("--max-pages", type=int, default=None, help="Limit number of pages per book (debug)")
    parser.add_argument(
        "--attn-impl",
        choices=["flash_attention_2", "sdpa", "eager", "auto"],
        default="auto",
        help="Attention backend for Qwen (flash_attention_2 recommended on T4 when available)",
    )
    parser.add_argument(
        "--load-in-8bit",
        action="store_true",
        help="Load the model with bitsandbytes 8-bit weights to fit smaller GPUs",
    )
    parser.add_argument(
        "--gpu-mem-limit",
        type=float,
        default=14.0,
        help="Per-GPU memory budget (GiB) for auto device_map when multiple GPUs exist",
    )
    parser.add_argument(
        "--cpu-mem-limit",
        type=float,
        default=29.0,
        help="CPU memory budget (GiB) for auto device_map spillover",
    )
    parser.add_argument("--log-level", default="INFO", help="Logging level")
    return parser.parse_args()


@dataclass
class AssetSpec:
    asset_id: str
    caption: Optional[str]
    bbox: Optional[Sequence[float]]
    asset_type: Optional[str]


class QwenOcrClient:
    """Wraps processor + model loading and inference for Qwen3-VL."""

    def __init__(
        self,
        model_id: str,
        device_map: str,
        torch_dtype: str,
        min_pixels: int,
        max_pixels: int,
        max_new_tokens: int,
        temperature: float,
        attn_impl: str,
        load_in_8bit: bool,
        gpu_mem_limit: float,
        cpu_mem_limit: float,
    ) -> None:
        dtype = self._resolve_dtype(torch_dtype)
        if dtype == "auto":
            dtype = self._default_dtype()
        LOGGER.info("Loading processor %s", model_id)
        self.processor = AutoProcessor.from_pretrained(
            model_id,
            trust_remote_code=True,
            min_pixels=min_pixels,
            max_pixels=max_pixels,
        )
        LOGGER.info("Loading model %s", model_id)
        model_kwargs: Dict[str, object] = {
            "trust_remote_code": True,
        }
        if attn_impl and attn_impl != "auto":
            model_kwargs["attn_implementation"] = attn_impl
        if load_in_8bit:
            try:
                from transformers import BitsAndBytesConfig  # type: ignore
            except ImportError as exc:  # noqa: F401
                raise RuntimeError(
                    "bitsandbytes is required for --load-in-8bit; install it via pip install bitsandbytes",
                ) from exc
            model_kwargs["quantization_config"] = BitsAndBytesConfig(load_in_8bit=True)
        else:
            model_kwargs["torch_dtype"] = dtype
        device_map_value = device_map or "auto"
        model_kwargs["device_map"] = device_map_value
        if device_map_value == "auto":
            max_memory = self._build_max_memory_dict(gpu_mem_limit, cpu_mem_limit)
            if max_memory:
                model_kwargs["max_memory"] = max_memory
        self.model = Qwen3VLForConditionalGeneration.from_pretrained(
            model_id,
            **model_kwargs,
        )
        self.model.eval()
        self.max_new_tokens = max_new_tokens
        self.temperature = max(0.0, float(temperature))
        self.do_sample = self.temperature > 0

    def _resolve_dtype(self, name: str):
        if name == "auto":
            return "auto"
        mapping = {
            "float16": torch.float16,
            "bfloat16": torch.bfloat16,
            "float32": torch.float32,
        }
        return mapping[name]

    def _default_dtype(self):
        if torch.cuda.is_available():
            return torch.float16
        return torch.float32

    def _build_max_memory_dict(self, gpu_limit: float, cpu_limit: float) -> Optional[Dict[object, str]]:
        if not torch.cuda.is_available():
            return None
        limits: Dict[object, str] = {}
        num_devices = torch.cuda.device_count()
        for idx in range(num_devices):
            total_gb = torch.cuda.get_device_properties(idx).total_memory / (1024**3)
            cap = min(max(total_gb - 1, 2.0), gpu_limit) if gpu_limit else total_gb - 1
            limits[idx] = f"{cap:.2f}GiB"
        if cpu_limit:
            limits["cpu"] = f"{cpu_limit:.0f}GiB"
        return limits

    def ocr_page(self, page_num: int, image_path: Path) -> Dict[str, object]:
        messages = [
            {"role": "system", "content": [{"type": "text", "text": PAGE_PROMPT}]},
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"Trang số {page_num:04d}. Hãy trả về JSON theo yêu cầu.",
                    },
                    {"type": "image", "image": str(image_path)},
                ],
            },
        ]
        prompt = self.processor.apply_chat_template(
            messages,
            add_generation_prompt=True,
            tokenize=False,
        )
        vision_inputs, _ = process_vision_info(messages)
        model_inputs = self.processor(
            text=[prompt],
            images=vision_inputs,
            return_tensors="pt",
        ).to(self.model.device)
        gen_kwargs = {
            "max_new_tokens": self.max_new_tokens,
            "do_sample": self.do_sample,
        }
        if self.do_sample:
            gen_kwargs["temperature"] = self.temperature
        generated = self.model.generate(
            **model_inputs,
            **gen_kwargs,
        )
        output = self.processor.batch_decode(generated, skip_special_tokens=True)[0]
        return self._parse_response(output)

    def _parse_response(self, output: str) -> Dict[str, object]:
        payload = extract_json_block(output)
        if payload is None:
            LOGGER.warning("Falling back to raw markdown because JSON parse failed")
            return {"markdown": output.strip(), "assets": []}
        return payload


class MarkdownDocxExporter:
    """Very small Markdown -> DOCX bridge for headings, lists, and figures."""

    def convert(self, markdown_text: str, markdown_path: Path, docx_path: Path) -> None:
        doc = Document()
        parent = markdown_path.parent
        for line in markdown_text.splitlines():
            stripped = line.rstrip()
            if not stripped:
                doc.add_paragraph("")
                continue
            if match := HEADING_RE.match(stripped):
                level = min(len(match.group(1)), 4)
                doc.add_heading(match.group(2).strip(), level=level)
                continue
            if match := BULLET_RE.match(stripped):
                doc.add_paragraph(match.group(1).strip(), style="List Bullet")
                continue
            if match := NUMBERED_RE.match(stripped):
                doc.add_paragraph(match.group(2).strip(), style="List Number")
                continue
            if match := IMAGE_LINE_RE.match(stripped):
                self._insert_image(doc, parent / match.group(2).strip(), match.group(1))
                continue
            if IMAGE_INLINE_RE.search(stripped):
                text_only = IMAGE_INLINE_RE.sub(lambda m: f"{m.group(1)} (xem hình {m.group(2)})", stripped)
                doc.add_paragraph(text_only)
                for img_match in IMAGE_INLINE_RE.finditer(stripped):
                    self._insert_image(doc, parent / img_match.group(2).strip(), img_match.group(1))
                continue
            doc.add_paragraph(stripped)
        doc.save(docx_path)

    def _insert_image(self, document: Document, image_path: Path, caption: str) -> None:
        try:
            paragraph = document.add_paragraph(caption or "Hình")
            paragraph.alignment = 1
            document.add_picture(str(image_path), width=Inches(5.5))
        except Exception as exc:  # noqa: BLE001
            LOGGER.warning("Không thể chèn hình %s: %s", image_path, exc)


class BookOcrPipeline:
    def __init__(self, args: argparse.Namespace) -> None:
        self.source_dir = args.source.resolve()
        self.output_dir = args.output.resolve()
        self.max_pages = args.max_pages
        self.ocr = QwenOcrClient(
            model_id=args.model_id,
            device_map=args.device_map,
            torch_dtype=args.dtype,
            min_pixels=args.min_pixels,
            max_pixels=args.max_pixels,
            max_new_tokens=args.max_new_tokens,
            temperature=args.temperature,
            attn_impl=args.attn_impl,
            load_in_8bit=args.load_in_8bit,
            gpu_mem_limit=args.gpu_mem_limit,
            cpu_mem_limit=args.cpu_mem_limit,
        )
        self.books = self._discover_books(args.books)
        self.docx_exporter = MarkdownDocxExporter()

    def _discover_books(self, requested: Optional[Sequence[str]]) -> List[str]:
        if requested:
            return list(requested)
        return sorted(
            [d.name for d in self.source_dir.iterdir() if d.is_dir()],
        )

    def run(self) -> None:
        for book in self.books:
            LOGGER.info("Processing %s", book)
            self._process_book(book)

    def _process_book(self, book: str) -> None:
        book_dir = self.source_dir / book
        page_meta_dir = book_dir / "pages_data"
        page_image_dir = book_dir / "pdf_images"
        crop_dir = book_dir / "crops"
        if not page_meta_dir.exists() or not page_image_dir.exists():
            LOGGER.warning("Book %s missing required folders", book)
            return
        output_root = self.output_dir / book
        markdown_dir = output_root / "markdown"
        assets_root = output_root / "assets"
        docx_dir = output_root / "docx"
        markdown_dir.mkdir(parents=True, exist_ok=True)
        assets_root.mkdir(parents=True, exist_ok=True)
        docx_dir.mkdir(parents=True, exist_ok=True)

        aggregated_sections: List[str] = []
        page_files = sorted(page_meta_dir.glob("*.json"))
        if self.max_pages:
            page_files = page_files[: self.max_pages]
        for json_path in tqdm(page_files, desc=f"{book} pages"):
            page_payload = json.loads(json_path.read_text())
            page_number = int(page_payload.get("page_num", 0))
            image_name = page_payload.get("source_image")
            if not image_name:
                LOGGER.warning("Skipping page %s with no image", json_path)
                continue
            image_path = page_image_dir / image_name
            if not image_path.exists():
                LOGGER.warning("Missing image %s", image_path)
                continue
            page_markdown, asset_summary = self._ocr_single_page(
                page_number=page_number,
                image_path=image_path,
                markdown_dir=markdown_dir,
                assets_root=assets_root,
                crop_dir=crop_dir,
            )
            aggregated_sections.append(page_markdown)
            page_md_path = markdown_dir / f"page_{page_number:04d}.md"
            page_md_path.write_text(page_markdown)
            LOGGER.debug("Wrote %s (%d assets)", page_md_path, asset_summary)
        if not aggregated_sections:
            LOGGER.warning("Book %s produced no pages", book)
            return
        book_markdown = "\n\n".join(aggregated_sections)
        book_md_path = markdown_dir / f"{book}.md"
        book_md_path.write_text(book_markdown)
        docx_path = docx_dir / f"{book}.docx"
        self.docx_exporter.convert(book_markdown, book_md_path, docx_path)
        LOGGER.info("Finished %s -> %s", book, docx_path)

    def _ocr_single_page(
        self,
        page_number: int,
        image_path: Path,
        markdown_dir: Path,
        assets_root: Path,
        crop_dir: Path,
    ) -> Tuple[str, int]:
        response = self.ocr.ocr_page(page_number, image_path)
        markdown_body = response.get("markdown", "").strip()
        raw_assets = response.get("assets", []) or []
        page_asset_dir = assets_root / f"page_{page_number:04d}"
        page_asset_dir.mkdir(parents=True, exist_ok=True)
        rendered_assets = self._materialize_assets(
            raw_assets=raw_assets,
            page_image=image_path,
            crop_dir=crop_dir,
            dest_dir=page_asset_dir,
        )
        markdown_with_assets = self._inject_assets(
            markdown_body,
            rendered_assets,
            markdown_dir=markdown_dir,
            page_number=page_number,
        )
        return markdown_with_assets.strip(), len(rendered_assets)

    def _materialize_assets(
        self,
        raw_assets: Iterable[Dict[str, object]],
        page_image: Path,
        crop_dir: Path,
        dest_dir: Path,
    ) -> Dict[str, Path]:
        image = Image.open(page_image)
        width, height = image.size
        rendered: Dict[str, Path] = {}
        try:
            for asset_dict in raw_assets:
                spec = AssetSpec(
                    asset_id=str(asset_dict.get("id") or asset_dict.get("asset_id") or f"asset-{len(rendered)+1}"),
                    caption=asset_dict.get("caption"),
                    bbox=asset_dict.get("bbox"),
                    asset_type=asset_dict.get("type"),
                )
                asset_path = dest_dir / f"asset_{slugify(spec.asset_id)}.png"
                bbox = normalize_bbox(spec.bbox)
                crop_img = None
                if bbox:
                    clamped = clamp_bbox_to_size(bbox, width, height)
                    if clamped:
                        crop_img = image.crop(clamped)
                elif ref := asset_dict.get("img_path"):
                    ref_path = crop_dir / ref
                    if ref_path.exists():
                        with Image.open(ref_path) as ref_img:
                            crop_img = ref_img.copy()
                if crop_img is None:
                    LOGGER.warning("Skipping asset %s (no bbox/img)", spec.asset_id)
                    continue
                crop_img.save(asset_path)
                crop_img.close()
                rendered[spec.asset_id] = asset_path
            return rendered
        finally:
            image.close()

    def _inject_assets(
        self,
        markdown: str,
        assets: Dict[str, Path],
        markdown_dir: Path,
        page_number: int,
    ) -> str:
        if not assets:
            return markdown

        used: set[str] = set()

        def replace(match: re.Match[str]) -> str:
            asset_id = match.group(1)
            if asset_id not in assets:
                return match.group(0)
            asset_path = assets[asset_id]
            rel_path = os.path.relpath(asset_path, markdown_dir)
            alt = asset_id.replace("-", " ")
            used.add(asset_id)
            return f"![{alt}]({rel_path})"

        rendered_markdown = ASSET_PLACEHOLDER_RE.sub(replace, markdown)
        leftover = [aid for aid in assets if aid not in used]
        if not leftover:
            return rendered_markdown
        lines = [rendered_markdown.rstrip(), "", "### Hình ảnh bổ sung"]
        for aid in leftover:
            rel_path = os.path.relpath(assets[aid], markdown_dir)
            lines.append(f"![{aid}]({rel_path})")
        return "\n".join(lines)


def normalize_bbox(bbox: Optional[Sequence[float]]) -> Optional[Tuple[int, int, int, int]]:
    if not bbox:
        return None
    if len(bbox) == 4 and all(isinstance(v, (int, float)) for v in bbox):
        x1, y1, x2, y2 = bbox
    elif isinstance(bbox[0], (list, tuple)):
        xs = [pt[0] for pt in bbox]
        ys = [pt[1] for pt in bbox]
        x1, y1, x2, y2 = min(xs), min(ys), max(xs), max(ys)
    else:
        return None
    return int(max(0, min(x1, x2))), int(max(0, min(y1, y2))), int(max(x1, x2)), int(max(y1, y2))


def clamp_bbox_to_size(bbox: Tuple[int, int, int, int], width: int, height: int) -> Optional[Tuple[int, int, int, int]]:
    left, top, right, bottom = bbox
    left = max(0, min(left, width))
    right = max(0, min(right, width))
    top = max(0, min(top, height))
    bottom = max(0, min(bottom, height))
    if right - left < 2 or bottom - top < 2:
        return None
    return int(left), int(top), int(right), int(bottom)


def slugify(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "-", value).strip("-").lower() or "asset"


def extract_json_block(text_blob: str) -> Optional[Dict[str, object]]:
    text_blob = text_blob.strip()
    candidates = []
    if "```" in text_blob:
        for block in re.findall(r"```(?:json)?\s*(.*?)```", text_blob, flags=re.S):
            candidates.append(block.strip())
    candidates.append(text_blob)
    decoder = json.JSONDecoder()
    for candidate in candidates:
        candidate = candidate.strip()
        if not candidate:
            continue
        try:
            parsed = json.loads(candidate)
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            pass
        idx = candidate.find("{")
        while idx != -1:
            try:
                parsed, _ = decoder.raw_decode(candidate[idx:])
            except json.JSONDecodeError:
                idx = candidate.find("{", idx + 1)
                continue
            if isinstance(parsed, dict):
                return parsed
            idx = candidate.find("{", idx + 1)
    return None


def configure_logging(level: str) -> None:
    logging.basicConfig(
        level=getattr(logging, level.upper(), logging.INFO),
        format="%(asctime)s - %(levelname)s - %(message)s",
    )


def main() -> None:
    args = parse_args()
    configure_logging(args.log_level)
    pipeline = BookOcrPipeline(args)
    pipeline.run()


if __name__ == "__main__":
    main()
