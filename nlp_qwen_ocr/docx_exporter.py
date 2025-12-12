"""Markdown to DOCX conversion helpers."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

IMAGE_LINE_RE = re.compile(r"^\[(.*?)\]\((.*?)\)$")
IMAGE_INLINE_RE = re.compile(r"\[(.*?)\]\((.*?)\)")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
NUMBERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")


class MarkdownDocxExporter:
    """Minimal Markdown -> DOCX renderer (headings, lists, figures)."""

    def convert(self, markdown_text: str, markdown_path: Path, docx_path: Path) -> None:
        doc = Document()
        self._ensure_vietnamese_fonts(doc)
        parent = markdown_path.parent
        for line in markdown_text.splitlines():
            stripped = line.rstrip()
            if not stripped:
                doc.add_paragraph("")
                continue
            if match := HEADING_RE.match(stripped):
                level = min(len(match.group(1)), 4)
                doc.add_heading(match.group(2).strip(), level=level)
                continue
            if match := BULLET_RE.match(stripped):
                doc.add_paragraph(match.group(1).strip(), style="List Bullet")
                continue
            if match := NUMBERED_RE.match(stripped):
                doc.add_paragraph(match.group(2).strip(), style="List Number")
                continue
            if match := IMAGE_LINE_RE.match(stripped):
                target = match.group(2).strip()
                if self._looks_like_asset(target):
                    caption = self._sanitize_label(match.group(1), target)
                    self._insert_image_link(doc, parent / target, caption, target)
                    continue
            inline_matches = []
            for img_match in IMAGE_INLINE_RE.finditer(stripped):
                target = img_match.group(2).strip()
                if not self._looks_like_asset(target):
                    continue
                caption = self._sanitize_label(img_match.group(1), target)
                inline_matches.append((img_match, target, caption))
            if inline_matches:
                text_only = stripped
                for img_match, target, caption in inline_matches:
                    readable_target = Path(target).name
                    snippet = f"{caption} (xem hình {readable_target})"
                    text_only = text_only.replace(img_match.group(0), snippet, 1)
                doc.add_paragraph(text_only)
                for img_match, target, caption in inline_matches:
                    self._insert_image_link(doc, parent / target, caption, target)
                continue
            doc.add_paragraph(stripped)
        doc.save(docx_path)

    def _insert_image_link(
        self,
        document: Document,
        image_path: Path,
        caption: str,
        markdown_target: str | None = None,
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = 1
        try:
            self._add_hyperlink(
                paragraph,
                markdown_target or image_path.name,
                image_path.resolve().as_uri(),
            )
        except Exception as exc:  # noqa: BLE001
            paragraph.add_run(f" [Không thể tạo liên kết: {exc}]")

    def _add_hyperlink(self, paragraph, text: str, url: str) -> None:
        """Attach a clickable hyperlink run to the given paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        r_style = OxmlElement("w:rStyle")
        r_style.set(qn("w:val"), "Hyperlink")
        r_pr.append(r_style)
        new_run.append(r_pr)

        text_elem = OxmlElement("w:t")
        text_elem.text = text
        new_run.append(text_elem)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _looks_like_asset(self, target: str) -> bool:
        cleaned = target.strip().lower()
        if not cleaned:
            return False
        return cleaned.endswith((".png", ".jpg", ".jpeg", ".webp")) and "/assets/" in cleaned

    def _sanitize_label(self, label: str | None, target: str) -> str:
        cleaned_label = (label or "").strip()
        if not cleaned_label:
            return "Hình minh họa"
        if self._label_looks_like_path(cleaned_label, target):
            return "Hình minh họa"
        return cleaned_label

    def _label_looks_like_path(self, label: str, target: str) -> bool:
        lowered = label.lower()
        target_path = target.lower()
        target_name = Path(target).name.lower()
        if lowered in {target_path, target_name}:
            return True
        if "/" in label or "\\" in label:
            return True
        return False

    def _ensure_vietnamese_fonts(self, document: Document) -> None:
        preferred_font = "Times New Roman"
        style_names = (
            "Normal",
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "Heading 4",
            "List Bullet",
            "List Number",
        )
        for style_name in style_names:
            try:
                style = document.styles[style_name]
            except KeyError:
                continue
            style.font.name = preferred_font
            self._set_style_font_face(style, preferred_font)

    def _set_style_font_face(self, style, font_name: str) -> None:
        get_or_add_rpr = getattr(style._element, "get_or_add_rPr", None)
        if get_or_add_rpr is not None:
            r_pr = get_or_add_rpr()
        else:
            r_pr = style._element._add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            add_r_fonts = getattr(r_pr, "get_or_add_rFonts", None)
            if add_r_fonts is not None:
                r_fonts = add_r_fonts()
            else:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            r_fonts.set(qn(attr), font_name)
